"""Generate the standard templates, the synthetic sample contracts
(.docx and .pdf), and the ground-truth files used by the eval.

Usage: .venv/bin/python src/generate_documents.py
"""
import copy
import sys
from pathlib import Path

import yaml
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

sys.path.insert(0, str(Path(__file__).parent))
import contract_texts as ct

ROOT = Path(__file__).parent.parent
TEMPLATES = ROOT / "templates"
SAMPLES = ROOT / "samples"
GROUND_TRUTH = SAMPLES / "ground_truth"


def build_clause_list(sample):
    """Apply a sample's replace/insert edits to its base template."""
    if sample["base"] == "nda":
        clauses = copy.deepcopy(ct.NDA_CLAUSES)
    elif sample["base"] == "order_form":
        clauses = copy.deepcopy(ct.OF_CLAUSES)
    else:
        return copy.deepcopy(ct.OFFER_LETTER_BODY)

    clauses = [(h, sample["replace"].get(h, t)) for h, t in clauses]
    for ins in sample.get("insert", []):
        idx = next(i for i, (h, _) in enumerate(clauses) if h == ins["after"])
        clauses.insert(idx + 1, (ins["heading"], ins["text"]))
    return clauses


def doc_meta(sample):
    if sample["base"] == "nda":
        title = ct.NDA_TITLE
        intro = ct.NDA_INTRO.format(date=sample["date"], party=sample["party"])
    elif sample["base"] == "order_form":
        title = ct.OF_TITLE
        intro = ct.OF_INTRO.format(
            date=sample["date"], party=sample["party"],
            msa_date=sample.get("msa_date", ""),
        )
    else:
        title = ct.OFFER_LETTER_TITLE
        intro = (f"Date: {sample['date']}. Private and confidential. "
                 f"Addressed to: {sample['party']}.")
    return title, intro


def clause_text(sample, clauses):
    """Substitute per-sample commercial details into clause bodies."""
    out = []
    for h, t in clauses:
        t = (t.replace("{seats}", sample.get("seats", ""))
              .replace("{fees}", sample.get("fees", "")))
        out.append((h, t))
    return out


def write_docx(path, title, intro, clauses, signature=True):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(10.5)

    h = doc.add_paragraph()
    run = h.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(intro)
    for i, (heading, text) in enumerate(clauses, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {heading}. ")
        run.bold = True
        p.add_run(text)
    if signature:
        doc.add_paragraph(
            "IN WITNESS WHEREOF, the parties have executed this document as of "
            "the date first written above.")
        doc.add_paragraph("Fictional Co Ltd\n\nBy: ____________________\n"
                          "Name:\nTitle:")
        doc.add_paragraph("Counterparty\n\nBy: ____________________\n"
                          "Name:\nTitle:")
    doc.save(path)


def write_pdf(path, title, intro, clauses):
    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["Normal"], fontName="Times-Roman",
                          fontSize=10.5, leading=14, spaceAfter=8)
    head = ParagraphStyle("head", parent=styles["Title"], fontName="Times-Bold",
                          fontSize=14)
    story = [Paragraph(title, head), Spacer(1, 10), Paragraph(intro, body)]
    for i, (heading, text) in enumerate(clauses, start=1):
        story.append(Paragraph(f"<b>{i}. {heading}.</b> {text}", body))
    story.append(Paragraph(
        "IN WITNESS WHEREOF, the parties have executed this document as of the "
        "date first written above.", body))
    SimpleDocTemplate(str(path), pagesize=letter).build(story)


def main():
    for d in (TEMPLATES, SAMPLES, GROUND_TRUTH):
        d.mkdir(parents=True, exist_ok=True)

    # Templates (the baseline the assistant diffs against)
    write_docx(TEMPLATES / "nda_template.docx", ct.NDA_TITLE,
               ct.NDA_INTRO.format(date="[Effective Date]", party="[Counterparty]"),
               ct.NDA_CLAUSES)
    of_clauses = [(h, t.replace("{seats}", "[Seats]").replace("{fees}", "[Fees]"))
                  for h, t in ct.OF_CLAUSES]
    write_docx(TEMPLATES / "order_form_template.docx", ct.OF_TITLE,
               ct.OF_INTRO.format(date="[Date]", party="[Customer]",
                                  msa_date="[MSA Date]"),
               of_clauses)

    for sample in ct.SAMPLES:
        clauses = clause_text(sample, build_clause_list(sample))
        title, intro = doc_meta(sample)
        fname = f"{sample['name']}.{sample['format']}"
        path = SAMPLES / fname
        if sample["format"] == "docx":
            write_docx(path, title, intro, clauses,
                       signature=sample["base"] != "offer_letter")
        else:
            write_pdf(path, title, intro, clauses)

        gt = dict(sample["ground_truth"])
        gt["file"] = fname
        with open(GROUND_TRUTH / f"{sample['name']}.yaml", "w") as f:
            yaml.safe_dump(gt, f, sort_keys=False)
        print(f"wrote {path}")

    print(f"wrote templates to {TEMPLATES}")


if __name__ == "__main__":
    main()
