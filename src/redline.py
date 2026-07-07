"""Produce a redlined .docx with real tracked changes (w:ins / w:del) and
margin comments from a verified analysis.

The redline is built on the original Word file so the lawyer keeps the
counterparty's formatting. For PDF input we synthesise a .docx from the
extracted text first (noted in the comments), since the deliverable the team
works in is Word.
"""
import difflib
import re
from datetime import datetime, timezone

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.run import Run
from docx.shared import Pt

AUTHOR = "Fictional Legal Triage Assistant"
W_DATE = datetime(2026, 7, 7, tzinfo=timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
_next_id = [1000]


def _wid():
    _next_id[0] += 1
    return str(_next_id[0])


def _norm(s):
    s = s.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")
    return re.sub(r"\s+", " ", s).strip()


def _para_char_map(p):
    """Return (full_text, bold_flags list) across the paragraph's runs."""
    text, bold = "", []
    for r in p.findall(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        is_bold = rpr is not None and rpr.find(qn("w:b")) is not None
        for t in r.findall(qn("w:t")):
            s = t.text or ""
            text += s
            bold.extend([is_bold] * len(s))
    return text, bold


def _find_span(haystack, needle):
    """Locate needle in haystack tolerating whitespace/smart-quote drift."""
    pattern = r"\s+".join(re.escape(tok) for tok in _norm(needle).split())
    pattern = (pattern.replace('\\"', '[“”"]')
                      .replace("\\'", "[‘’']"))
    m = re.search(pattern, haystack)
    return (m.start(), m.end()) if m else None


def _make_run(text, bold=False, deleted=False):
    r = OxmlElement("w:r")
    if bold:
        rpr = OxmlElement("w:rPr")
        rpr.append(OxmlElement("w:b"))
        r.append(rpr)
    t = OxmlElement("w:delText" if deleted else "w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _tracked(tag, runs):
    el = OxmlElement(tag)
    el.set(qn("w:id"), _wid())
    el.set(qn("w:author"), AUTHOR)
    el.set(qn("w:date"), W_DATE)
    for r in runs:
        el.append(r)
    return el


def _plain_runs(text, bold_flags):
    """Emit runs for unchanged text, splitting at bold boundaries."""
    out = []
    i = 0
    while i < len(text):
        j = i
        while j < len(text) and bold_flags[j] == bold_flags[i]:
            j += 1
        out.append(_make_run(text[i:j], bold=bold_flags[i]))
        i = j
    return out


def _diff_elements(old, new):
    """Word-level diff between old and new text -> list of elements."""
    tokens_old = re.findall(r"\S+\s*", old)
    tokens_new = re.findall(r"\S+\s*", new)
    sm = difflib.SequenceMatcher(a=tokens_old, b=tokens_new, autojunk=False)
    elements = []
    for op, a0, a1, b0, b1 in sm.get_opcodes():
        old_seg, new_seg = "".join(tokens_old[a0:a1]), "".join(tokens_new[b0:b1])
        if op == "equal":
            elements.append(_make_run(old_seg))
        if op in ("delete", "replace") and old_seg:
            elements.append(_tracked("w:del", [_make_run(old_seg, deleted=True)]))
        if op in ("insert", "replace") and new_seg:
            elements.append(_tracked("w:ins", [_make_run(new_seg)]))
    return elements


def _rebuild(p, elements):
    """Replace a paragraph's content (keeping pPr) with new elements."""
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    for el in elements:
        p.append(el)


def _anchor_comment(doc, para, text):
    """Attach a margin comment covering the whole paragraph."""
    rs = para._p.findall(f".//{qn('w:r')}")
    if not rs:
        return
    runs = [Run(rs[0], para)] if len(rs) == 1 else [Run(rs[0], para), Run(rs[-1], para)]
    doc.add_comment(runs=runs, text=text, author=AUTHOR, initials="FL")
    # add_comment may nest range markers inside w:ins/w:del wrappers; Word
    # expects them as direct children of w:p - hoist any that are nested.
    p = para._p
    for tag in ("w:commentRangeStart", "w:commentRangeEnd", "w:commentReference"):
        for el in p.findall(f".//{qn(tag)}"):
            parent = el.getparent()
            holder = parent if parent.tag == qn("w:r") and tag == "w:commentReference" else el
            container = holder.getparent()
            if container.tag != qn("w:p"):
                container.addnext(holder) if tag != "w:commentRangeStart" \
                    else container.addprevious(holder)
    # w:commentRangeEnd must precede the w:commentReference run.
    for ref in p.findall(qn("w:r")):
        if ref.find(qn("w:commentReference")) is None:
            continue
        cid = ref.find(qn("w:commentReference")).get(qn("w:id"))
        for end in p.findall(qn("w:commentRangeEnd")):
            if end.get(qn("w:id")) == cid and p.index(end) > p.index(ref):
                ref.addprevious(end)


def _mark_paragraph_deleted(p):
    """Mark the paragraph mark itself deleted so accepting merges cleanly."""
    ppr = p.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        p.insert(0, ppr)
    rpr = ppr.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        ppr.append(rpr)
    d = OxmlElement("w:del")
    d.set(qn("w:id"), _wid())
    d.set(qn("w:author"), AUTHOR)
    d.set(qn("w:date"), W_DATE)
    rpr.append(d)


def synthesize_docx_from_paragraphs(paragraphs, out_path):
    """For PDF input: build a Word document from the extracted text."""
    doc = Document()
    doc.styles["Normal"].font.name = "Georgia"
    doc.styles["Normal"].font.size = Pt(10.5)
    for i, text in enumerate(paragraphs):
        p = doc.add_paragraph()
        if i == 0:
            run = p.add_run(text)
            run.bold = True
        else:
            m = re.match(r"^(\d{1,2}\.\s+[A-Z][A-Za-z \-]{0,40}?\.)\s*(.*)$", text, re.S)
            if m:
                p.add_run(m.group(1) + " ").bold = True
                p.add_run(m.group(2))
            else:
                p.add_run(text)
    doc.save(out_path)
    return out_path


def apply_redlines(source_docx, analysis, out_path):
    """Apply tracked changes + comments for each verified deviation."""
    doc = Document(str(source_docx))
    paras = doc.paragraphs
    applied, skipped = [], []

    for dev in analysis.get("deviations", []):
        target = _locate_paragraph(paras, dev)
        if target is None:
            skipped.append(dev["clause_heading"])
            continue
        action = dev.get("action", "flag_only")
        comment = dev.get("comment", "")
        if dev.get("rule_id"):
            comment = f"[{dev['rule_id']}] {comment}"

        if action in ("replace", "delete") and dev.get("verified"):
            ok = _apply_text_change(target, dev)
            if not ok:
                action = "flag_only"
                comment = "[could not anchor change - review manually] " + comment
        else:
            action = "flag_only" if action in ("replace", "delete") else action

        _anchor_comment(doc, target, comment)
        applied.append((dev["clause_heading"], action))

    doc.save(str(out_path))
    return applied, skipped


def _locate_paragraph(paras, dev):
    snippet = dev.get("original_snippet") or ""
    idx = dev.get("paragraph_index")
    # The analysis paragraph indices count non-empty paragraphs.
    non_empty = [p for p in paras if p.text.strip()]
    if isinstance(idx, int) and 0 <= idx < len(non_empty):
        cand = non_empty[idx]
        if not snippet or _find_span(_norm(cand.text), _norm(snippet)):
            return cand
    if snippet:
        for p in non_empty:
            if _find_span(_norm(p.text), _norm(snippet)):
                return p
    heading = _norm(dev.get("clause_heading") or "")
    if heading:
        for p in non_empty:
            if heading.lower() in _norm(p.text).lower()[:120]:
                return p
    return None


def _apply_text_change(para, dev):
    p = para._p
    full, bold = _para_char_map(p)
    span = _find_span(full, dev["original_snippet"])
    if span is None:
        return False
    s, e = span
    replacement = dev.get("proposed_text") or ""
    if dev["action"] == "delete" and not replacement.strip():
        # Strike the entire clause, heading included.
        elements = [_tracked("w:del", _to_deleted(_plain_runs(full, bold)))]
        _rebuild(p, elements)
        _mark_paragraph_deleted(p)
        return True

    elements = []
    if s > 0:
        elements.extend(_plain_runs(full[:s], bold[:s]))
    elements.extend(_diff_elements(full[s:e], replacement))
    if e < len(full):
        elements.extend(_plain_runs(full[e:], bold[e:]))
    _rebuild(p, elements)
    return True


def _to_deleted(runs):
    for r in runs:
        for t in r.findall(qn("w:t")):
            t.tag = qn("w:delText")
    return runs
